"""
Parses resumes and job descriptions from PDF, DOCX, or TXT into raw text.
Output feeds into chunker.py.
"""
import os
from pathlib import Path
from dataclasses import dataclass

from pypdf import PdfReader
from docx import Document


@dataclass
class ParsedDocument:
    doc_id: str          # filename without extension
    doc_type: str         # "resume" | "jd"
    source_path: str
    raw_text: str


def parse_pdf(path: str) -> str:
    reader = PdfReader(path)
    text = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text.append(page_text)
    return "\n".join(text)


def parse_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def parse_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_file(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(path)
    elif ext == ".docx":
        return parse_docx(path)
    elif ext == ".txt":
        return parse_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def parse_directory(dir_path: str, doc_type: str) -> list[ParsedDocument]:
    """Parse all supported files in a directory into ParsedDocument objects."""
    results = []
    supported = {".pdf", ".docx", ".txt"}

    for fname in sorted(os.listdir(dir_path)):
        fpath = os.path.join(dir_path, fname)
        ext = Path(fname).suffix.lower()
        if ext not in supported or not os.path.isfile(fpath):
            continue

        try:
            raw_text = parse_file(fpath)
        except Exception as e:
            print(f"[parser] Failed to parse {fname}: {e}")
            continue

        if not raw_text.strip():
            print(f"[parser] Warning: empty text extracted from {fname}")
            continue

        results.append(ParsedDocument(
            doc_id=Path(fname).stem,
            doc_type=doc_type,
            source_path=fpath,
            raw_text=raw_text,
        ))

    return results


if __name__ == "__main__":
    import argparse
    from app.config import RESUME_DIR, JD_DIR

    ap = argparse.ArgumentParser()
    ap.add_argument("--input", choices=["resumes", "jds", "both"], default="both")
    args = ap.parse_args()

    if args.input in ("resumes", "both"):
        resumes = parse_directory(RESUME_DIR, "resume")
        print(f"Parsed {len(resumes)} resume(s) from {RESUME_DIR}")

    if args.input in ("jds", "both"):
        jds = parse_directory(JD_DIR, "jd")
        print(f"Parsed {len(jds)} JD(s) from {JD_DIR}")
